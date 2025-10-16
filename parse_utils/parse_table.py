from docx import Document
import pandas as pd


def parse_first_table(path):
    doc = Document(path)
    table = doc.tables[0]

    data = {}
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if not any(cells) or cells[0] in ['General information:', '', 'Report date:']:
            continue
        key = cells[0]
        value_candidates = cells[2:]
        value = next((v for v in value_candidates if v), '')
        data[key] = value

    df = pd.DataFrame([data])

    df = remove_columns(df)
    return rename_columns(df)

def parse_analysis_table(table):
    """
    Парсит таблицу анализа, объединяет Parameter + Method + Unit, удаляет Contract/U,
    и транспонирует результат (ключи становятся колонками, Result — строками)
    """
    table_data = []
    found_header = False

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if not any(cells):
            continue
        if not found_header and cells[0].lower().startswith("parameter"):
            headers = cells
            found_header = True
            continue
        if found_header:
            if len(cells) == len(headers):
                table_data.append(cells)

    df = pd.DataFrame(table_data, columns=headers)

    # Объединяем столбцы
    columns_to_combine = ['Parameter', 'Method', 'Unit']
    existing_columns = [col for col in columns_to_combine if col in df.columns]
    df['Combined'] = df[existing_columns].fillna('').agg(' / '.join, axis=1)
    # Оставляем только нужные столбцы
    df = df[['Combined', 'Result']]
    # преобразовывает данные в float
    value_cols = [i for i, col in enumerate(df.columns) if col == 'Result']
    for i in value_cols:
        df.iloc[:, i] = df.iloc[:, i].apply(convert_to_float)
    # транспонируем результат
    transposed = df.set_index('Combined').T

    return transposed.reset_index(drop=True)


def parse_second_table(filepath: str) -> pd.DataFrame:
    doc = Document(filepath)
    df = pd.DataFrame()
    targets = ['Approved by Not approved',
               'Approved by Head of Laboratory',
               'Checked by Technician laboratory assistant',
               'Checked by Laboratory assistant',
               'Checked by Deputy Chief of the Chromatography Department'
               'Weighted average',
               'Checked by Deputy Chief of the Basic Chemistry Department',
               'Conclusion',
               'Name of Pesticide']

    for table in doc.tables[1:]:
        skip_table = False
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(x in cells for x in targets):
                skip_table = True
                break
        if skip_table:
            continue  # Переход к следующей таблице

        # Таблица — анализы
        analysis_df = parse_analysis_table(table)

        # Объединяем по горизонтали
        df = pd.concat([df, analysis_df], axis=1)

        df = remove_columns(df)

    return df


def rename_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Переименовывает колонки DataFrame согласно заданной карте соответствий.
    """
    column_map = {
        'Lab#': 'Laboratory Number',
        'Seals': 'Seal',
        'Weight of sample:': 'Weight of sample',
        'Quantity represented by this sample, mt:': 'Quantity by sample',
    }

    return df.rename(columns=column_map, errors='ignore')


def remove_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Удаляет колонки из DataFrame.
    """
    columns_to_remove = [
        'EXTRA INFO / EXTRA INFO / EXTRA INFO',
        'Sample description (as labeled):',
        'Instructions received',
        'Sampling',
        'Date of sampling',
        'Dates of testing',
        'Date of Analyses start',
        'Date of Analyses finish'
    ]
    return df.drop(columns=columns_to_remove, errors='ignore')

def convert_to_float(val):
    try:
        return float(val)
    except (ValueError, TypeError):
        return val  # оставим как есть, если не число
